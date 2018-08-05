#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Fri Aug  3 17:42:09 2018

@author: root
"""
import os
import json
import argparse
from docx import Document
from docx.shared import Inches

parser = argparse.ArgumentParser('Tabulate the result in a docx file')
parser.add_argument('-m','--master_path',help='The folder that contains many model folders')
parser.add_argument('-d','--doc',help='path to docx file, /path/result.docx')
args = parser.parse_args()

master_path = args.master_path
doc_path = args.doc

#Add each row of information of one model
def add_row(config_dict, result_dict, table, count, folder_path):
    count_dict = {'Model':str(count)}
    combine_dict = {**count_dict,**config_dict, **result_dict}
    
    cells = table.add_row().cells
    for count, key in enumerate(combine_dict):
        if key != 'gif':
            cells[count].text = str(combine_dict[key])
        else:
            cell_img = cells[count].paragraphs[0]
            run = cell_img.add_run()
            run.add_picture(os.path.join(folder_path,combine_dict[key]), width=Inches(1.0))


def tabulate(master_path):
    folder_1 = os.listdir(master_path)[0]
    folder_1_path = os.path.join(master_path,folder_1)    
    count_dict = {'Model':''}
    
    config_dict = json.loads(open(os.path.join(folder_1_path,'config.json'),'r').read())
    result_dict = json.loads(open(os.path.join(folder_1_path,'result.json'),'r').read())
    
    combine_dict = {**count_dict, **config_dict, **result_dict}
    
    
    num_col = len(combine_dict)
    document = Document()
    
    document.add_paragraph('Results')
    #Create all the headings first
    table = document.add_table(rows=1, cols=num_col)
    heading_cells = table.rows[0].cells
    for count, key in enumerate(combine_dict):
        heading_cells[count].text = key
    
    for folder in os.listdir(master_path):
        count = folder.split('_')[1]
        folder_path = os.path.join(master_path, folder)
        config_dict = json.loads(open(os.path.join(folder_path,'config.json'),'r').read())
        result_dict = json.loads(open(os.path.join(folder_path,'result.json'),'r').read())
        add_row(config_dict,result_dict,table,count,folder_path)
    document.save(doc_path)

tabulate(master_path)